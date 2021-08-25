import pyttsx3
import speech_recognition as sr
import docx


engine = pyttsx3.init('sapi5')
voices = engine.getProperty('voices')
engine.setProperty('voice' , voices[0].id)
r = sr.Recognizer()


def speak(audio):
    engine.say(audio)
    engine.runAndWait()
    print("this is me")
    print("hello")



def newornot():
    with sr.Microphone() as source:
        speak("if this is a new assingment say yes else say no")
        print("lisening yes or no....")
        r.pause_threshold = 1
        audio = r.listen(source)
    new_or_not = r.recognize_google(audio , language='en-in')
    if "yes" in new_or_not :
        yes = "yes"
        return yes 
        
    elif "no" in new_or_not :
        speak("specify the name of the assignment file")
        with sr.Microphone() as source:
            print("listening for the name")
            r.pause_threshold = 1
            audio = r.listen(source)
            name =  r.recognize_google(audio , language='en-in')
            name = name + ".docx"
            return name 


    
argument = newornot()
noargument = argument
 

#creates a new doc or saves to old doc according to new or not fucnt

if argument == "yes":
    speak("pls provide a name for the new assignment")
    #defines a new doc name
    with sr.Microphone() as source:
        print("listening for the name")
        r.pause_threshold = 1
        audio = r.listen(source)
        name_new =  r.recognize_google(audio , language='en-in')
        name_new = name_new + ".docx"
    doc = docx.Document()
    
else : 
    try: 
        
        doc = docx.Document(noargument)
    except : 
        print("enter a valid name say it without the extention")
        speak("enter a valid name say it without the extention")
        


def listener():

    with sr.Microphone() as source:
        speak("say what you want to write")
        print("Listening.......")
        r.pause_threshold = 1
        audio = r.listen(source)

    try : 
        print("recognizing......")
        user = r.recognize_google(audio , language='en-in')
        print(f"user said: {user}")
        doc.add_paragraph(user)
        if argument == "yes":
            doc.save(name_new)
        else:
            doc.save(noargument)
    except : 
        print("repeat that pls")
        speak("repeat that pls")
        
    return user

def text2speech():
    text = listener()
    speak(text)

    
    
if __name__ == '__main__':
    speak("hello user i am the bot")
    text2speech()
    
    
